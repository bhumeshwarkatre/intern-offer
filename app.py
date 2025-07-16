import os
import re
import csv
import qrcode
import random
import string
import tempfile
import base64
import streamlit as st
from datetime import date
from smtplib import SMTP
from docxtpl import DocxTemplate
from docx.shared import Inches
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email import encoders
import pandas as pd

# ✅ Aspose Cloud
import asposewordscloud
from asposewordscloud.apis.words_api import WordsApi
from asposewordscloud.models.requests import UploadFileRequest, SaveAsRequest, DownloadFileRequest
from asposewordscloud.models import PdfSaveOptionsData

# ✅ Google Sheets
from google.oauth2.service_account import Credentials
import gspread

# --- Config ---
st.set_page_config("Intern Offer Generator", layout="wide")
EMAIL = st.secrets["email"]["user"]
PASSWORD = st.secrets["email"]["password"]
ADMIN_KEY = st.secrets["admin"]["key"]
CSV_FILE = "intern_offers.csv"
TEMPLATE_FILE = os.path.join(tempfile.gettempdir(), "offer_template.docx")
LOGO = "logo.png"
STORAGE_NAME = "demo"  # ✅ Your Aspose storage name

# --- Aspose Setup ---
api_sid = st.secrets["aspose"]["app_sid"]
api_key = st.secrets["aspose"]["app_key"]
words_api = WordsApi(api_sid, api_key)

# --- Google Sheets Setup ---
def get_gsheet():
    scope = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive"
    ]
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=scope
    )
    client = gspread.authorize(creds)
    return client.open("intern_offers").sheet1

def save_to_gsheet(data):
    sheet = get_gsheet()
    row = [data['i_id'], data['intern_name'], data['domain'], data['start_date'], data['end_date'], data['offer_date'], data['email']]
    sheet.append_row(row)

# --- Template base64 decode ---
if not os.path.exists(TEMPLATE_FILE):
    encoded_template = st.secrets["template_base64"]["template_base64"]
    with open(TEMPLATE_FILE, "wb") as f:
        f.write(base64.b64decode(encoded_template))

# --- Styling ---
st.markdown("""
<style>
    .title-text {
        font-size: 2rem;
        font-weight: 700;
    }
    .stButton>button {
        background-color: #1E88E5;
        color: white;
        padding: 0.5rem 1.5rem;
        border-radius: 8px;
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

# --- Header ---
with st.container():
    col_logo, col_title = st.columns([1, 6])
    with col_logo:
        if os.path.exists(LOGO):
            st.image(LOGO, width=80)
    with col_title:
        st.markdown('<div class="title-text">SkyHighes Technologies Internship Letter Portal</div>', unsafe_allow_html=True)

st.divider()

# --- Utility Functions ---
def format_date(d):
    return d.strftime("%A, %d %B %Y")

def generate_certificate_key():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=9))

def generate_qr(data):
    qr = qrcode.QRCode(box_size=10, border=0)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    path = os.path.join(tempfile.gettempdir(), "qr.png")
    img.save(path)
    return path

def save_to_csv(data):
    exists = os.path.exists(CSV_FILE)
    with open(CSV_FILE, mode='a', newline='') as f:
        writer = csv.writer(f)
        if not exists:
            writer.writerow(["Intern ID", "Intern Name", "Domain", "Start Date", "End Date", "Offer Date", "Email"])
        writer.writerow([data['i_id'], data['intern_name'], data['domain'], data['start_date'], data['end_date'], data['offer_date'], data['email']])

def send_email(receiver, pdf_path, data):
    msg = MIMEMultipart()
    msg['From'] = EMAIL
    msg['To'] = receiver
    msg['Subject'] = f"🎉 Congratulations {data['intern_name']}! Your Internship Offer"

    html = f"""
    <html>
    <body style="font-family: Arial, sans-serif; background-color: #f9f9f9; padding: 20px;">
      <div style="background-color: #fff; padding: 20px; border-radius: 10px;">
        <p>Dear {data["intern_name"]},</p>
        <p>We are delighted to offer you an <strong>Internship Opportunity</strong> at <strong>SkyHighes Technology</strong>! 🎉</p>
        <h3>Internship Details:</h3>
        <ul>
          <li><b>Intern Name:</b> {data["intern_name"]}</li>
          <li><b>Domain:</b> {data["domain"]}</li>
          <li><b>Start Date:</b> {data["start_date"]}</li>
          <li><b>End Date:</b> {data["end_date"]}</li>
          <li><b>Offer Date:</b> {data["offer_date"]}</li>
        </ul>
        <p>Your offer letter is attached as a PDF document. Kindly review and confirm your acceptance.</p>
        <p>We look forward to working with you!</p>
        <br>
        <p><strong>SkyHighes Technology Team</strong></p>
      </div>
    </body>
    </html>
    """
    msg.attach(MIMEText(html, 'html'))

    with open(pdf_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        filename = os.path.basename(pdf_path)
        part.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(part)

    with SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(EMAIL, PASSWORD)
        server.send_message(msg)

# --- Form ---
with st.form("offer_form"):
    st.subheader("🎓 Internship Offer Letter Generator")

    col1, col2, col3 = st.columns(3)
    with col1:
        intern_name = st.text_input("Intern Name")
    with col2:
        domain = st.text_input("Domain")
    with col3:
        email = st.text_input("Recipient Email")

    col4, col5, col6 = st.columns(3)
    with col4:
        start_date = st.date_input("Start Date", value=date.today())
    with col5:
        end_date = st.date_input("End Date", value=date.today())
    with col6:
        offer_date = st.date_input("Offer Date", value=date.today())

    submit = st.form_submit_button("🚀 Generate & Send Offer Letter")

# --- On Submit ---
if submit:
    if not all([intern_name, domain, email]):
        st.error("❌ Please fill all fields.")
    elif not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        st.warning("⚠️ Invalid email.")
    elif end_date < start_date:
        st.warning("⚠️ End date cannot be before start date.")
    else:
        intern_id = generate_certificate_key()
        data = {
            "intern_name": intern_name.title().strip(),
            "domain": domain.title().strip(),
            "start_date": format_date(start_date),
            "end_date": format_date(end_date),
            "offer_date": format_date(offer_date),
            "i_id": intern_id,
            "email": email.lower().strip()
        }

        save_to_csv(data)
        try:
            save_to_gsheet(data)
        except Exception as e:
            st.warning(f"⚠️ Google Sheet sync failed: {e}")

        doc = DocxTemplate(TEMPLATE_FILE)
        doc.render(data)

        qr_path = generate_qr(f"{intern_name}, {domain}, {start_date}, {end_date}, {offer_date}, {intern_id}")
        try:
            doc.tables[0].rows[0].cells[2].paragraphs[0].add_run().add_picture(qr_path, width=Inches(1.4))
        except:
            st.warning("⚠️ QR insertion failed.")

        docx_path = os.path.join(tempfile.gettempdir(), f"Offer_{intern_name}.docx")
        doc.save(docx_path)

        cloud_doc_name = f"{intern_id}.docx"
        cloud_pdf_name = f"Offer_{intern_name}.pdf"
        local_pdf_path = os.path.join(tempfile.gettempdir(), cloud_pdf_name)

        try:
            with open(docx_path, "rb") as f:
                upload_request = UploadFileRequest(
                    file_content=f,
                    path=cloud_doc_name,
                    storage_name=STORAGE_NAME,
                    folder="/"
                )
                upload_result = words_api.upload_file(upload_request)

            if not upload_result.uploaded or cloud_doc_name not in upload_result.uploaded:
                st.error(f"❌ Upload to Aspose failed. File {cloud_doc_name} not found.")
            else:
                save_opts = PdfSaveOptionsData(file_name=cloud_pdf_name)
                save_request = SaveAsRequest(
                    name=cloud_doc_name,
                    save_options_data=save_opts,
                    storage_name=STORAGE_NAME,
                    folder="/"
                )
                save_result = words_api.save_as(save_request)
                cloud_pdf_path = save_result.save_result.dest_document.href

                pdf_stream = words_api.download_file(
                    DownloadFileRequest(cloud_pdf_path, storage_name=STORAGE_NAME)
                )
                with open(local_pdf_path, "wb") as f:
                    f.write(pdf_stream)

                send_email(email, local_pdf_path, data)
                st.success(f"✅ Offer letter sent to {email}")

                with open(local_pdf_path, "rb") as f:
                    st.download_button("📥 Download Offer Letter", f, file_name=os.path.basename(local_pdf_path))

        except Exception as e:
            st.error(f"❌ Error occurred: {e}")

# --- Admin Panel ---
st.divider()
with st.expander("🔐 Admin Panel"):
    admin_key = st.text_input("Enter Admin Key", type="password")
    if admin_key == ADMIN_KEY:
        st.success("✅ Access granted.")
        if os.path.exists(CSV_FILE):
            try:
                df = pd.read_csv(CSV_FILE)
                if not df.empty:
                    st.dataframe(df)
                    with open(CSV_FILE, "rb") as f_dl:
                        st.download_button("📥 Download CSV", f_dl, file_name="intern_offers.csv")
                else:
                    st.info("CSV file is empty.")
            except Exception as e:
                st.error(f"Error reading CSV: {e}")
        else:
            st.info("CSV log not found.")

        st.divider()
        st.markdown("<h3 style='color:#1E88E5;'>📥 One-Time CSV Upload <small style='color:gray;'>(Optional)</small></h3>", unsafe_allow_html=True)
        uploaded_csv = st.file_uploader("Upload Existing Intern CSV", type=["csv"])
        if uploaded_csv is not None:
            try:
                with open(CSV_FILE, "wb") as f:
                    f.write(uploaded_csv.read())
                st.success("✅ Uploaded and saved CSV successfully.")
            except Exception as e:
                st.error(f"Error saving uploaded CSV: {e}")
    elif admin_key:
        st.error("❌ Invalid key.")

# 🔽 Footer
st.markdown("<hr><center><small>© 2025 SkyHighes Technologies. All Rights Reserved.</small></center>", unsafe_allow_html=True)
